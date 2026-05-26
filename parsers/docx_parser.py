from docx import Document as DocxDocument
from pathlib import Path
from parsers.base_parser import BaseParser, Document

class WordParser(BaseParser):
    def parse(self, file_path: str) -> Document:
        doc = DocxDocument(file_path)
        content = self.extract_text(file_path)
        return Document(
            content=content,
            metadata={
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            },
            format="docx",
            pages=[content]
        )

    def extract_text(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return "\n".join(texts)
