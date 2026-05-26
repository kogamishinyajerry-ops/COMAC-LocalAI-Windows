"""
COMAC 报告生成器 - 多格式报告输出

支持格式:
- Excel (带国企标准样式)
- Word (DOCX)
- HTML (可预览)
- Markdown
"""
from typing import List, Dict, Optional, Any
from dataclasses import dataclass, field
from pathlib import Path
from datetime import datetime
import html


@dataclass
class ReportSection:
    """报告章节"""
    title: str
    content: str
    level: int = 2
    items: List[Dict] = field(default_factory=list)
    table_data: Optional[List[Dict]] = None
    table_headers: Optional[List[str]] = None


@dataclass
class Report:
    """报告定义"""
    title: str
    subtitle: str = ""
    author: str = "COMAC AI"
    date: str = ""
    sections: List[ReportSection] = field(default_factory=list)
    tags: List[str] = field(default_factory=list)


def _h(text: str) -> str:
    """HTML 内容安全转义（防止 XSS）"""
    if text is None:
        return ""
    return html.escape(str(text))


class ReportGenerator:
    """
    报告生成器

    使用方式:
        gen = ReportGenerator()
        report = gen.create_report("审查报告", sections=[...])
        
        # 生成 Word
        gen.to_docx(report, "output/report.docx")
        
        # 生成 Markdown
        md = gen.to_markdown(report)
    """

    def __init__(self, output_dir: str = "temp/outputs"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def create_report(
        self,
        title: str,
        sections: List[ReportSection],
        subtitle: str = "",
        author: str = "COMAC AI",
        tags: List[str] = None,
    ) -> Report:
        """创建报告对象"""
        return Report(
            title=title,
            subtitle=subtitle,
            author=author,
            date=datetime.now().strftime("%Y-%m-%d %H:%M"),
            sections=sections,
            tags=tags or [],
        )

    def to_markdown(self, report: Report) -> str:
        """导出为 Markdown"""
        lines = []

        # 标题
        lines.append(f"# {report.title}")
        if report.subtitle:
            lines.append(f"*{report.subtitle}*")
        lines.append("")
        lines.append(f"**编制**: {report.author} | **日期**: {report.date}")
        if report.tags:
            lines.append(f"**标签**: {', '.join(report.tags)}")
        lines.append("")
        lines.append("---")
        lines.append("")

        # 章节
        for section in report.sections:
            prefix = "#" * min(section.level + 1, 6)
            lines.append(f"{prefix} {section.title}")
            lines.append("")

            if section.content:
                lines.append(section.content)
                lines.append("")

            if section.items:
                for item in section.items:
                    lines.append(f"- **{item.get('title', '')}**: {item.get('content', '')}")
                lines.append("")

            if section.table_data and section.table_headers:
                # Markdown 表格
                headers = section.table_headers
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row in section.table_data:
                    cells = [str(row.get(h, "")) for h in headers]
                    lines.append("| " + " | ".join(cells) + " |")
                lines.append("")

        # 页脚
        lines.append("---")
        lines.append(f"*本报告由 COMAC AI 文档处理平台自动生成*")

        return "\n".join(lines)

    def to_html(self, report: Report, template: str = "default") -> str:
        """导出为 HTML"""
        md_content = self.to_markdown(report)

        if template == "comac":
            return self._comac_html_template(report)
        elif template == "minimal":
            return f"<pre style='white-space:pre-wrap;font-family:monospace'>{md_content}</pre>"
        else:
            # 默认: 带样式的 HTML
            return self._default_html_template(report)

    def _default_html_template(self, report: Report) -> str:
        """默认 HTML 模板（XSS 安全版）"""
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{_h(report.title)}</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; color: #333; }}
        h1 {{ color: #1A3A5C; border-bottom: 3px solid #1A3A5C; padding-bottom: 10px; }}
        h2 {{ color: #2A5A8C; margin-top: 30px; }}
        h3 {{ color: #3A6A9C; }}
        .meta {{ color: #666; font-size: 14px; margin-bottom: 20px; }}
        .section {{ margin-bottom: 30px; }}
        .content {{ line-height: 1.8; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th {{ background: #1A3A5C; color: white; padding: 10px; text-align: left; }}
        td {{ border: 1px solid #ddd; padding: 8px; }}
        tr:nth-child(even) {{ background: #f9f9f9; }}
        .tags {{ margin: 10px 0; }}
        .tag {{ display: inline-block; background: #E8F5E9; color: #2E7D32; padding: 3px 8px; border-radius: 3px; margin: 2px; font-size: 12px; }}
        footer {{ margin-top: 50px; padding-top: 20px; border-top: 1px solid #ddd; color: #999; font-size: 12px; }}
    </style>
</head>
<body>
    <h1>{_h(report.title)}</h1>
    {f'<p class="meta"><strong>{_h(report.subtitle)}</strong></p>' if report.subtitle else ''}
    <p class="meta">编制: {_h(report.author)} | 日期: {_h(report.date)}</p>
    {self._render_tags_html(report.tags)}
"""
        for section in report.sections:
            level = section.level
            tag = f"h{min(level + 1, 6)}"
            html += f"<div class='section'><{tag}>{_h(section.title)}</{tag}>"

            if section.content:
                html += f"<div class='content'>{_h(section.content)}</div>"

            if section.items:
                html += "<ul>"
                for item in section.items:
                    html += f"<li><strong>{_h(item.get('title', ''))}</strong>: {_h(item.get('content', ''))}</li>"
                html += "</ul>"

            if section.table_data and section.table_headers:
                html += "<table><tr>"
                for h in section.table_headers:
                    html += f"<th>{_h(h)}</th>"
                html += "</tr>"
                for row in section.table_data:
                    html += "<tr>"
                    for h in section.table_headers:
                        html += f"<td>{_h(row.get(h, ''))}</td>"
                    html += "</tr>"
                html += "</table>"

            html += "</div>"

        html += f"""
    <footer>本报告由 COMAC AI 文档处理平台自动生成 | {_h(report.date)}</footer>
</body>
</html>"""
        return html

    def _comac_html_template(self, report: Report) -> str:
        """COMAC 品牌 HTML 模板 (国企标准, XSS 安全版)"""
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{_h(report.title)}</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Microsoft YaHei', sans-serif; background: #F0F2F5; }}
        .header {{ background: linear-gradient(135deg, #1A3A5C 0%, #2A5A8C 100%); color: white; padding: 40px; text-align: center; }}
        .header h1 {{ font-size: 24px; margin-bottom: 8px; }}
        .header p {{ font-size: 14px; opacity: 0.85; }}
        .container {{ max-width: 960px; margin: 0 auto; padding: 20px 40px; }}
        .meta-card {{ background: white; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .meta-card .row {{ display: flex; gap: 30px; color: #666; font-size: 14px; }}
        .section-card {{ background: white; border-radius: 8px; padding: 25px; margin: 20px 0; box-shadow: 0 2px 8px rgba(0,0,0,0.08); }}
        .section-card h2 {{ color: #1A3A5C; font-size: 18px; margin-bottom: 15px; padding-bottom: 10px; border-bottom: 2px solid #E8F5E9; }}
        .section-card h3 {{ color: #2A5A8C; font-size: 16px; margin: 15px 0 10px; }}
        .content {{ line-height: 1.8; color: #333; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 14px; }}
        th {{ background: #1A3A5C; color: white; padding: 12px 10px; text-align: left; font-weight: 600; }}
        td {{ border: 1px solid #E0E0E0; padding: 10px; }}
        tr:nth-child(even) {{ background: #FAFAFA; }}
        .tag {{ display: inline-block; background: #E8F5E9; color: #2E7D32; padding: 4px 10px; border-radius: 4px; margin: 3px; font-size: 12px; }}
        .conclusion {{ background: #FFF8E1; border-left: 4px solid #FFC107; padding: 15px 20px; margin: 20px 0; border-radius: 4px; }}
        .stamp {{ text-align: right; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #999; font-size: 12px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{_h(report.title)}</h1>
        {f'<p>{_h(report.subtitle)}</p>' if report.subtitle else ''}
    </div>
    <div class="container">
        <div class="meta-card">
            <div class="row">
                <span>📋 编制单位: {_h(report.author)}</span>
                <span>📅 日期: {_h(report.date)}</span>
                <span>🔒 密级: 内部</span>
            </div>
            {self._render_tags_html(report.tags)}
        </div>
"""
        for section in report.sections:
            html += f"""
        <div class="section-card">
            <h2>{_h(section.title)}</h2>"""

            if section.content:
                html += f"<div class='content'>{_h(section.content)}</div>"

            if section.items:
                html += "<ul style='padding-left:20px;line-height:2'>"
                for item in section.items:
                    title = item.get('title', '')
                    content = item.get('content', '')
                    if title:
                        html += f"<li><strong>{_h(title)}</strong>: {_h(content)}</li>"
                    else:
                        html += f"<li>{_h(content)}</li>"
                html += "</ul>"

            if section.table_data and section.table_headers:
                html += "<table><thead><tr>"
                for h in section.table_headers:
                    html += f"<th>{_h(h)}</th>"
                html += "</tr></thead><tbody>"
                for row in section.table_data:
                    html += "<tr>"
                    for h in section.table_headers:
                        html += f"<td>{_h(row.get(h, ''))}</td>"
                    html += "</tr>"
                html += "</tbody></table>"

            html += "</div>"

        html += f"""
        <div class="stamp">
            COMAC AI 文档处理平台 | 自动生成于 {_h(report.date)}
        </div>
    </div>
</body>
</html>"""
        return html

    def _render_tags_html(self, tags: List[str]) -> str:
        if not tags:
            return ""
        tag_html = " ".join([f'<span class="tag">{t}</span>' for t in tags])
        return f'<div class="tags">{tag_html}</div>'

    def to_docx(self, report: Report, filepath: str = None) -> str:
        """导出为 Word 文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("需要 python-docx 库")

        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)

        # 标题
        title_para = doc.add_heading(report.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if report.subtitle:
            sub = doc.add_paragraph(report.subtitle)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.runs[0].font.size = Pt(12)

        # 元信息
        meta = doc.add_paragraph(f"编制: {report.author} | 日期: {report.date} | 密级: 内部")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in meta.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        # 章节
        for section in report.sections:
            doc.add_heading(section.title, level=section.level)

            if section.content:
                p = doc.add_paragraph(section.content)

            if section.items:
                for item in section.items:
                    p = doc.add_paragraph(style='List Bullet')
                    title = item.get('title', '')
                    content = item.get('content', '')
                    if title:
                        p.add_run(f"{title}: ").bold = True
                    p.add_run(content)

            if section.table_data and section.table_headers:
                table = doc.add_table(rows=1, cols=len(section.table_headers))
                table.style = 'Light Grid Accent 1'

                # 表头
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(section.table_headers):
                    hdr_cells[i].text = h
                    for para in hdr_cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True

                # 数据行
                for row_data in section.table_data:
                    row_cells = table.add_row().cells
                    for i, h in enumerate(section.table_headers):
                        row_cells[i].text = str(row_data.get(h, ""))

                doc.add_paragraph()

        # 文件路径
        if filepath is None:
            filepath = str(self.output_dir / f"{report.title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        doc.save(filepath)
        return filepath

    def to_text(self, report: Report) -> str:
        """导出为纯文本"""
        lines = [f"{'=' * 60}", f"  {report.title}"]
        if report.subtitle:
            lines.append(f"  {report.subtitle}")
        lines.extend([f"{'=' * 60}", "", f"编制: {report.author}  日期: {report.date}", ""])

        for section in report.sections:
            lines.append(f"\n{'─' * 50}")
            lines.append(f"  {section.title}")
            lines.append(f"{'─' * 50}")
            if section.content:
                lines.append(section.content)
            if section.items:
                for item in section.items:
                    lines.append(f"  • {item.get('title', '')}: {item.get('content', '')}")

        lines.append(f"\n{'=' * 60}")
        lines.append(f"  COMAC AI 文档处理平台自动生成")
        lines.append(f"{'=' * 60}")
        return "\n".join(lines)
